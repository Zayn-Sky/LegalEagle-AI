from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path
from typing import Any, Dict, Iterable, List

from docx import Document as DocxDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
from tqdm import tqdm

PROJECT_ROOT = Path(__file__).resolve().parents[1]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from legaleagle_agent.settings import REQUIRED_CATEGORY_CODES

STATUS_LABELS = {
    1: "已废止",
    2: "已修改",
    3: "有效",
    4: "尚未生效",
}

ARTICLE_RE = re.compile(r"^(第[一二三四五六七八九十百千万零〇两\d]+条)\s*(.*)")
HIERARCHY_RE = re.compile(r"^(第[一二三四五六七八九十百千万零〇两\d]+[编章节])\s*(.*)|^(附则)$")
DROP_LINE_RE = re.compile(r"^[-—－\s]*\d+[-—－\s]*$")


def load_json(path: Path) -> Any:
    return json.loads(path.read_text(encoding="utf-8"))


def resolve_input_path(value: str, raw_dir: Path) -> Path:
    path = Path(value)
    if path.is_absolute() or path.exists():
        return path
    return raw_dir / path


def iter_jsonl(path: Path) -> Iterable[Dict[str, Any]]:
    if not path.exists():
        return
    with path.open("r", encoding="utf-8") as file:
        for line in file:
            line = line.strip()
            if line:
                yield json.loads(line)


def write_jsonl(path: Path, rows: Iterable[Dict[str, Any]]) -> int:
    path.parent.mkdir(parents=True, exist_ok=True)
    count = 0
    with path.open("w", encoding="utf-8") as file:
        for row in rows:
            file.write(json.dumps(row, ensure_ascii=False) + "\n")
            count += 1
    return count


def normalize_line(text: str) -> str:
    text = text.replace("\u3000", " ")
    text = re.sub(r"\s+", " ", text).strip()
    return text


def extract_docx_lines(path: Path) -> List[str]:
    doc = DocxDocument(path)
    lines: List[str] = []
    for paragraph in doc.paragraphs:
        text = normalize_line(paragraph.text)
        if text and not DROP_LINE_RE.match(text):
            lines.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [normalize_line(cell.text) for cell in row.cells]
            line = " ".join(cell for cell in cells if cell)
            if line and not DROP_LINE_RE.match(line):
                lines.append(line)
    return lines


def split_units(lines: List[str]) -> List[Dict[str, Any]]:
    units: List[Dict[str, Any]] = []
    hierarchy_state: Dict[str, str | None] = {"part": None, "chapter": None, "section": None, "appendix": None}
    front_matter: List[str] = []
    current: Dict[str, Any] | None = None

    def current_path() -> List[str]:
        if hierarchy_state["appendix"]:
            return [hierarchy_state["appendix"]]
        return [
            value
            for value in (
                hierarchy_state["part"],
                hierarchy_state["chapter"],
                hierarchy_state["section"],
            )
            if value
        ]

    def flush_current() -> None:
        nonlocal current
        if current and current["lines"]:
            current["text"] = "\n".join(current.pop("lines"))
            units.append(current)
        current = None

    for line in lines:
        article_match = ARTICLE_RE.match(line)
        hierarchy_match = HIERARCHY_RE.match(line)

        if article_match:
            flush_current()
            article = article_match.group(1)
            rest = article_match.group(2).strip()
            current = {
                "unit_type": "article",
                "article": article,
                "section_path": current_path(),
                "lines": [line if rest else article],
            }
            continue

        if hierarchy_match:
            flush_current()
            heading = line
            level_marker = hierarchy_match.group(1) or hierarchy_match.group(3) or heading
            if level_marker.endswith("编"):
                hierarchy_state.update({"part": heading, "chapter": None, "section": None, "appendix": None})
            elif level_marker.endswith("章"):
                hierarchy_state.update({"chapter": heading, "section": None, "appendix": None})
            elif level_marker.endswith("节"):
                hierarchy_state.update({"section": heading, "appendix": None})
            elif level_marker == "附则":
                hierarchy_state.update({"part": None, "chapter": None, "section": None, "appendix": heading})
            units.append({"unit_type": "heading", "article": "", "section_path": current_path(), "text": heading})
            continue

        if current:
            current["lines"].append(line)
        else:
            front_matter.append(line)

    flush_current()
    if front_matter:
        units.insert(0, {"unit_type": "front_matter", "article": "", "section_path": [], "text": "\n".join(front_matter)})
    return units


def build_splitter(chunk_size: int, chunk_overlap: int) -> RecursiveCharacterTextSplitter:
    return RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=[
            "\n\n",
            "\n第",
            "\n",
            "。", "；", "，",
            " ", "",
        ],
    )


def law_metadata(record: Dict[str, Any], docx_path: Path) -> Dict[str, Any]:
    detail = record["detail"]
    list_record = record.get("list_record", {})
    bbbs = detail.get("bbbs") or list_record.get("bbbs")
    return {
        "bbbs": bbbs,
        "title": detail.get("title") or list_record.get("title"),
        "category_code": record.get("category_code"),
        "category_name": record.get("category_name"),
        "status_code": detail.get("sxx") or list_record.get("sxx"),
        "status": STATUS_LABELS.get(detail.get("sxx") or list_record.get("sxx")),
        "law_type": detail.get("flxz") or list_record.get("flxz"),
        "issuing_authority": detail.get("zdjgName") or list_record.get("zdjgName"),
        "publish_date": detail.get("gbrq") or list_record.get("gbrq"),
        "effective_date": detail.get("sxrq") or list_record.get("sxrq"),
        "source": record.get("source"),
        "source_url": record.get("source_url"),
        "detail_url": record.get("detail_url"),
        "docx_path": str(docx_path),
    }


def validate_manifest(manifest: Dict[str, Any], raw_dir: Path) -> None:
    records = manifest.get("records") or []
    if not records:
        raise RuntimeError("manifest.json contains no records.")

    category_codes = {int(item["code"]) for item in manifest.get("categories", []) if item.get("code") is not None}
    missing_categories = REQUIRED_CATEGORY_CODES - category_codes
    if missing_categories:
        raise RuntimeError(f"manifest.json is missing required categories: {sorted(missing_categories)}")

    expected_count = sum(int(item.get("count", 0)) for item in manifest.get("categories", []))
    if expected_count != len(records):
        raise RuntimeError(
            f"manifest record count mismatch: categories expect {expected_count}, records contain {len(records)}"
        )

    seen = set()
    for item in records:
        bbbs = item.get("bbbs")
        if not bbbs:
            raise RuntimeError(f"manifest record has no bbbs: {item}")
        if bbbs in seen:
            raise RuntimeError(f"duplicate manifest record: {bbbs}")
        seen.add(bbbs)

        record_path = resolve_input_path(item["record_path"], raw_dir)
        docx_path = resolve_input_path(item["docx_path"], raw_dir)
        if not record_path.exists():
            raise FileNotFoundError(f"Missing record JSON: {record_path}")
        if not docx_path.exists():
            raise FileNotFoundError(f"Missing DOCX: {docx_path}")


def render_chunk_text(meta: Dict[str, Any], unit: Dict[str, Any], text: str) -> str:
    path = " > ".join(unit.get("section_path") or [])
    lines = [
        f"法律名称：{meta.get('title', '')}",
        f"法律部门：{meta.get('category_name', '')}",
        f"时效性：{meta.get('status', '')}",
        f"制定机关：{meta.get('issuing_authority', '')}",
        f"公布日期：{meta.get('publish_date', '')}",
        f"施行日期：{meta.get('effective_date', '')}",
    ]
    if path:
        lines.append(f"位置：{path}")
    if unit.get("article"):
        lines.append(f"条款：{unit['article']}")
    lines.append("正文：")
    lines.append(text)
    return "\n".join(lines)


def preprocess(args: argparse.Namespace) -> None:
    raw_dir = Path(args.raw_dir)
    output_dir = Path(args.output_dir)
    manifest_path = raw_dir / "manifest.json"
    if not manifest_path.exists():
        raise FileNotFoundError(f"Missing required manifest: {manifest_path}")
    manifest = load_json(manifest_path)
    validate_manifest(manifest, raw_dir)
    splitter = build_splitter(args.chunk_size, args.chunk_overlap)

    laws: List[Dict[str, Any]] = []
    chunks: List[Dict[str, Any]] = []

    for item in tqdm(manifest["records"], desc="preprocess"):
        record_path = resolve_input_path(item["record_path"], raw_dir)
        docx_path = resolve_input_path(item["docx_path"], raw_dir)

        record = load_json(record_path)
        meta = law_metadata(record, docx_path)
        lines = extract_docx_lines(docx_path)
        full_text = "\n".join(lines)
        units = split_units(lines)
        laws.append({**meta, "text": full_text, "unit_count": len(units)})

        for unit_index, unit in enumerate(units):
            unit_text = unit["text"].strip()
            if not unit_text:
                continue
            pieces = splitter.split_text(unit_text) if len(unit_text) > args.chunk_size else [unit_text]
            for piece_index, piece in enumerate(pieces):
                chunk_id = f"{meta['bbbs']}:{unit_index}:{piece_index}"
                chunk_meta = {
                    **meta,
                    "chunk_id": chunk_id,
                    "unit_index": unit_index,
                    "piece_index": piece_index,
                    "unit_type": unit.get("unit_type"),
                    "article": unit.get("article") or "",
                    "section_path": " > ".join(unit.get("section_path") or []),
                }
                chunks.append(
                    {
                        "id": chunk_id,
                        "text": render_chunk_text(meta, unit, piece),
                        "metadata": chunk_meta,
                    }
                )

    law_count = write_jsonl(output_dir / "laws.jsonl", laws)
    chunk_count = write_jsonl(output_dir / "chunks.jsonl", chunks)
    print(f"Wrote {law_count} laws and {chunk_count} chunks to {output_dir}")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Preprocess NPC DOCX laws into RAG JSONL chunks.")
    parser.add_argument("--raw-dir", default="data/raw/npc")
    parser.add_argument("--output-dir", default="data/processed")
    parser.add_argument("--chunk-size", type=int, default=1000)
    parser.add_argument("--chunk-overlap", type=int, default=160)
    return parser.parse_args()


if __name__ == "__main__":
    preprocess(parse_args())
